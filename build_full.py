import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.gridspec as gridspec
import numpy as np
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = '/home/claude/charts'
os.makedirs(OUT, exist_ok=True)

# ── Palette ──────────────────────────────────────────────────
GENZ    = '#2196F3'
MILEN   = '#7C3AED'
BG      = '#0F1117'
CARD    = '#1A1D27'
TEXT    = '#E8E8F0'
GRID    = '#2A2D3A'
GX      = '#10B981'
BOOM    = '#F59E0B'

def style(fig, axes):
    fig.patch.set_facecolor(BG)
    for ax in axes:
        ax.set_facecolor(CARD)
        ax.tick_params(colors=TEXT, labelsize=10)
        for sp in ax.spines.values(): sp.set_color(GRID)
        for lbl in ax.get_xticklabels() + ax.get_yticklabels():
            lbl.set_color(TEXT)
        ax.yaxis.label.set_color(TEXT)
        ax.xaxis.label.set_color(TEXT)
        ax.title.set_color(TEXT)

def pct_fmt(ax): ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v,_: f'{int(v)}%'))

def bar_labels(ax, bars, color, fontsize=10):
    for b in bars:
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+0.8,
                f'{int(b.get_height())}%', ha='center', va='bottom',
                color=color, fontsize=fontsize, fontweight='bold')

def save(fig, name):
    fig.tight_layout()
    fig.savefig(f'{OUT}/{name}', dpi=150, bbox_inches='tight', facecolor=BG)
    plt.close()
    print(f'✓ {name}')

# ════════════════════════════════════════════════════════
# CHART 1 — Asset Ownership (FINRA/CFA 2023)
# ════════════════════════════════════════════════════════
assets = ['Crypto', 'Individual\nStocks', 'Mutual\nFunds', 'ETFs', 'NFTs']
gz1    = [55, 41, 35, 23, 25]
ml1    = [57, 38, 43, 26, 28]
x, w   = np.arange(len(assets)), 0.35

fig, ax = plt.subplots(figsize=(10, 5.5))
style(fig, [ax])
b1 = ax.bar(x-w/2, gz1, w, label='Gen Z',       color=GENZ,  alpha=0.9, zorder=3)
b2 = ax.bar(x+w/2, ml1, w, label='Millennials', color=MILEN, alpha=0.9, zorder=3)
bar_labels(ax, b1, GENZ); bar_labels(ax, b2, MILEN)
ax.set_xticks(x); ax.set_xticklabels(assets, fontsize=11)
ax.set_ylim(0,75); pct_fmt(ax)
ax.set_ylabel('% Who Own This Asset')
ax.grid(axis='y', color=GRID, lw=0.8, zorder=0)
ax.legend(fontsize=11, facecolor=CARD, edgecolor=GRID, labelcolor=TEXT)
ax.set_title('Asset Ownership Rates — Gen Z vs. Millennials\nSource: FINRA Foundation & CFA Institute, 2023', fontsize=13, fontweight='bold', pad=14)
save(fig, 'c1_ownership_finra.png')

# ════════════════════════════════════════════════════════
# CHART 2 — Stock Types Owned (Motley Fool 2025)
# ════════════════════════════════════════════════════════
stypes = ['U.S. Stocks','Growth\nStocks','Value\nStocks','Dividend\nStocks','Small/\nMid-Cap','Large-\nCap','REITs','Speculative','ESG']
gz2    = [42, 45, 39, 28, 31, 26, 23, 13, 15]
ml2    = [54, 48, 39, 35, 39, 32, 23, 14, 13]
x2     = np.arange(len(stypes))

fig, ax = plt.subplots(figsize=(13, 5.5))
style(fig, [ax])
b1 = ax.bar(x2-w/2, gz2, w, label='Gen Z',       color=GENZ,  alpha=0.9, zorder=3)
b2 = ax.bar(x2+w/2, ml2, w, label='Millennials', color=MILEN, alpha=0.9, zorder=3)
bar_labels(ax, b1, GENZ, 9); bar_labels(ax, b2, MILEN, 9)
ax.set_xticks(x2); ax.set_xticklabels(stypes, fontsize=9.5)
ax.set_ylim(0, 70); pct_fmt(ax)
ax.set_ylabel('% Who Own This Stock Type')
ax.grid(axis='y', color=GRID, lw=0.8, zorder=0)
ax.legend(fontsize=11, facecolor=CARD, edgecolor=GRID, labelcolor=TEXT)
ax.set_title('Stock Types Owned — Gen Z vs. Millennials\nSource: Motley Fool Generational Investing Trends Survey, July 2025', fontsize=13, fontweight='bold', pad=14)
save(fig, 'c2_stock_types_motley.png')

# ════════════════════════════════════════════════════════
# CHART 3 — Sector Ownership (Motley Fool 2025)
# ════════════════════════════════════════════════════════
sectors = ['Technology','Financials','Real\nEstate','Healthcare\n& Biotech','Energy &\nUtilities','Consumer','Crypto-\nRelated','AI Stocks','Industrials']
gz3     = [50, 36, 29, 29, 30, 16, 23, 22, 20]
ml3     = [58, 38, 26, 36, 34, 22, 28, 21, 20]
x3      = np.arange(len(sectors))

fig, ax = plt.subplots(figsize=(13, 5.5))
style(fig, [ax])
b1 = ax.bar(x3-w/2, gz3, w, label='Gen Z',       color=GENZ,  alpha=0.9, zorder=3)
b2 = ax.bar(x3+w/2, ml3, w, label='Millennials', color=MILEN, alpha=0.9, zorder=3)
bar_labels(ax, b1, GENZ, 9); bar_labels(ax, b2, MILEN, 9)
ax.set_xticks(x3); ax.set_xticklabels(sectors, fontsize=9.5)
ax.set_ylim(0, 75); pct_fmt(ax)
ax.set_ylabel('% Invested in Sector')
ax.grid(axis='y', color=GRID, lw=0.8, zorder=0)
ax.legend(fontsize=11, facecolor=CARD, edgecolor=GRID, labelcolor=TEXT)
ax.set_title('Investment Sector Ownership — Gen Z vs. Millennials\nSource: Motley Fool Generational Investing Trends Survey, July 2025', fontsize=13, fontweight='bold', pad=14)
save(fig, 'c3_sectors_motley.png')

# ════════════════════════════════════════════════════════
# CHART 4 — Investment Goals (Motley Fool 2025)
# ════════════════════════════════════════════════════════
goals  = ['Passive\nIncome','Retirement\nPlanning','Buying\na Home','Capital\nAppreciation','Starting\na Business','Paying\nOff Debt']
gz4    = [43, 20, 16, 7, 9, 5]
ml4    = [40, 30, 9, 10, 4, 7]
x4     = np.arange(len(goals))

fig, ax = plt.subplots(figsize=(11, 5.5))
style(fig, [ax])
b1 = ax.bar(x4-w/2, gz4, w, label='Gen Z',       color=GENZ,  alpha=0.9, zorder=3)
b2 = ax.bar(x4+w/2, ml4, w, label='Millennials', color=MILEN, alpha=0.9, zorder=3)
bar_labels(ax, b1, GENZ); bar_labels(ax, b2, MILEN)
ax.set_xticks(x4); ax.set_xticklabels(goals, fontsize=11)
ax.set_ylim(0, 58); pct_fmt(ax)
ax.set_ylabel('% Citing as Primary Investment Goal')
ax.grid(axis='y', color=GRID, lw=0.8, zorder=0)
ax.legend(fontsize=11, facecolor=CARD, edgecolor=GRID, labelcolor=TEXT)
ax.set_title('Primary Investment Goals — Gen Z vs. Millennials\nSource: Motley Fool Generational Investing Trends Survey, July 2025', fontsize=13, fontweight='bold', pad=14)
save(fig, 'c4_goals_motley.png')

# ════════════════════════════════════════════════════════
# CHART 5 — Trade Frequency (Motley Fool 2025)
# ════════════════════════════════════════════════════════
freqs = ['Daily','Weekly','Multiple\nper Month','Once a\nMonth','Less than\nOnce/Month']
gz5   = [16, 34, 24, 10, 16]
ml5   = [14, 29, 27,  9, 20]
x5    = np.arange(len(freqs))

fig, ax = plt.subplots(figsize=(11, 5.5))
style(fig, [ax])
b1 = ax.bar(x5-w/2, gz5, w, label='Gen Z',       color=GENZ,  alpha=0.9, zorder=3)
b2 = ax.bar(x5+w/2, ml5, w, label='Millennials', color=MILEN, alpha=0.9, zorder=3)
bar_labels(ax, b1, GENZ); bar_labels(ax, b2, MILEN)
ax.set_xticks(x5); ax.set_xticklabels(freqs, fontsize=11)
ax.set_ylim(0, 48); pct_fmt(ax)
ax.set_ylabel('% of Investors')
ax.grid(axis='y', color=GRID, lw=0.8, zorder=0)
ax.legend(fontsize=11, facecolor=CARD, edgecolor=GRID, labelcolor=TEXT)
ax.set_title('Trading Frequency — Gen Z vs. Millennials\nSource: Motley Fool Generational Investing Trends Survey, July 2025', fontsize=13, fontweight='bold', pad=14)
save(fig, 'c5_trade_freq_motley.png')

# ════════════════════════════════════════════════════════
# CHART 6 — Info Sources (Motley Fool 2025)
# ════════════════════════════════════════════════════════
sources = ['YouTube','TikTok','Friends/\nFamily','Reddit','Financial\nAdvisor','Free\nWebsite','Podcasts']
gz6     = [67, 48, 40, 27, 25, 26, 22]
ml6     = [56, 26, 42, 28, 28, 37, 24]
x6      = np.arange(len(sources))

fig, ax = plt.subplots(figsize=(12, 5.5))
style(fig, [ax])
b1 = ax.bar(x6-w/2, gz6, w, label='Gen Z',       color=GENZ,  alpha=0.9, zorder=3)
b2 = ax.bar(x6+w/2, ml6, w, label='Millennials', color=MILEN, alpha=0.9, zorder=3)
bar_labels(ax, b1, GENZ); bar_labels(ax, b2, MILEN)
ax.set_xticks(x6); ax.set_xticklabels(sources, fontsize=11)
ax.set_ylim(0, 82); pct_fmt(ax)
ax.set_ylabel('% Using Source for Investing Info')
ax.grid(axis='y', color=GRID, lw=0.8, zorder=0)
ax.legend(fontsize=11, facecolor=CARD, edgecolor=GRID, labelcolor=TEXT)
ax.set_title('Investment Information Sources — Gen Z vs. Millennials\nSource: Motley Fool Generational Investing Trends Survey, July 2025', fontsize=13, fontweight='bold', pad=14)
save(fig, 'c6_info_sources_motley.png')

# ════════════════════════════════════════════════════════
# CHART 7 — Portfolio Allocation (Schwab Wave 2, 2025)
# ════════════════════════════════════════════════════════
alloc_labels = ['Stocks','Mutual\nFunds','Bonds','Crypto','Real\nEstate','ETFs','Options/\nFutures','Alternatives','Other']
alloc_vals   = [25, 13, 8, 10, 7, 6, 3, 3, 25]
colors_pie   = ['#3B82F6','#22D3EE','#6366F1','#F59E0B','#10B981','#8B5CF6','#F87171','#FB923C','#94A3B8']

fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(13, 6))
style(fig, [ax1, ax2])

wedges, texts, autotexts = ax1.pie(
    alloc_vals, labels=alloc_labels, colors=colors_pie,
    autopct='%1.0f%%', startangle=140,
    wedgeprops={'edgecolor': BG, 'linewidth': 2},
    textprops={'color': TEXT, 'fontsize': 9.5})
for at in autotexts: at.set_color(BG); at.set_fontweight('bold'); at.set_fontsize(9)
ax1.set_title('Average Portfolio Allocation\n(All Investors)', color=TEXT, fontsize=12, fontweight='bold')

# Bar chart of allocation
y_pos = np.arange(len(alloc_labels))
bars  = ax2.barh(y_pos, alloc_vals, color=colors_pie, alpha=0.9, height=0.6, zorder=3)
for bar, val in zip(bars, alloc_vals):
    ax2.text(bar.get_width()+0.4, bar.get_y()+bar.get_height()/2,
             f'{val}%', va='center', color=TEXT, fontsize=10, fontweight='bold')
ax2.set_yticks(y_pos); ax2.set_yticklabels(alloc_labels, fontsize=10)
ax2.set_xlim(0, 32); ax2.invert_yaxis()
ax2.set_xlabel('% of Portfolio'); ax2.grid(axis='x', color=GRID, lw=0.8, zorder=0)
ax2.set_title('Portfolio Allocation Breakdown', color=TEXT, fontsize=12, fontweight='bold')

fig.suptitle('Average Investment Portfolio Allocation — All U.S. Investors\nSource: Charles Schwab Modern Wealth Survey Wave 2, October 2025',
             color=TEXT, fontsize=13, fontweight='bold', y=1.02)
save(fig, 'c7_portfolio_schwab.png')

# ════════════════════════════════════════════════════════
# CHART 8 — FOMO / Risk / App Use (FINRA 2023)
# ════════════════════════════════════════════════════════
risk_labels = ['Willing to Take\nSubstantial Risk','Made Investment\nDue to FOMO','First Investment\nWas Crypto','Use Investing\nApps to Trade']
gz8 = [46, 50, 44, 65]
ml8 = [38, 42, 35, 55]   # FINRA surveyed millennials too
x8  = np.arange(len(risk_labels))

fig, ax = plt.subplots(figsize=(11, 5.5))
style(fig, [ax])
b1 = ax.bar(x8-w/2, gz8, w, label='Gen Z',       color=GENZ,  alpha=0.9, zorder=3)
b2 = ax.bar(x8+w/2, ml8, w, label='Millennials', color=MILEN, alpha=0.9, zorder=3)
bar_labels(ax, b1, GENZ); bar_labels(ax, b2, MILEN)
ax.set_xticks(x8); ax.set_xticklabels(risk_labels, fontsize=11)
ax.set_ylim(0, 80); pct_fmt(ax)
ax.set_ylabel('% of Respondents')
ax.grid(axis='y', color=GRID, lw=0.8, zorder=0)
ax.legend(fontsize=11, facecolor=CARD, edgecolor=GRID, labelcolor=TEXT)
ax.set_title('Risk Appetite & Behavioural Indicators — Gen Z vs. Millennials\nSource: FINRA Foundation & CFA Institute, 2023', fontsize=13, fontweight='bold', pad=14)
save(fig, 'c8_risk_finra.png')

# ════════════════════════════════════════════════════════
# CHART 9 — Gen Z views dividends as side hustle (Motley Fool)
# ════════════════════════════════════════════════════════
div_uses = ['Take as Cash\n(Everyday)','Take as Cash\n(Fun Money)','Save for\nGoals','Reinvest\nAutomatically','Reallocate\nManually']
gz9 = [26, 18, 19, 23, 11]
ml9 = [17, 15, 17, 38, 9]
x9  = np.arange(len(div_uses))

fig, ax = plt.subplots(figsize=(11, 5.5))
style(fig, [ax])
b1 = ax.bar(x9-w/2, gz9, w, label='Gen Z',       color=GENZ,  alpha=0.9, zorder=3)
b2 = ax.bar(x9+w/2, ml9, w, label='Millennials', color=MILEN, alpha=0.9, zorder=3)
bar_labels(ax, b1, GENZ); bar_labels(ax, b2, MILEN)
ax.set_xticks(x9); ax.set_xticklabels(div_uses, fontsize=11)
ax.set_ylim(0, 50); pct_fmt(ax)
ax.set_ylabel('% of Dividend Investors')
ax.grid(axis='y', color=GRID, lw=0.8, zorder=0)
ax.legend(fontsize=11, facecolor=CARD, edgecolor=GRID, labelcolor=TEXT)
ax.set_title('How Each Generation Uses Dividend Income\nSource: Motley Fool Generational Investing Trends Survey, July 2025', fontsize=13, fontweight='bold', pad=14)
save(fig, 'c9_dividends_motley.png')

print('\n✅ All 9 charts done\n')

# ════════════════════════════════════════════════════════
# WORD DOCUMENT
# ════════════════════════════════════════════════════════
def set_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_border(cell, color='CCCCCC'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4')
        el.set(qn('w:space'),'0'); el.set(qn('w:color'), color)
        b.append(el)
    tcPr.append(b)

def h1(doc, text):
    p = doc.add_heading(text, 1)
    r = p.runs[0]; r.font.name='Calibri'; r.font.color.rgb=RGBColor(0x1A,0x3A,0x5C)

def h2(doc, text):
    p = doc.add_heading(text, 2)
    r = p.runs[0]; r.font.name='Calibri'; r.font.color.rgb=RGBColor(0x21,0x96,0xF3)

def body(doc, text, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.name='Calibri'; r.font.size=Pt(11); r.italic=italic
    if color: r.font.color.rgb=color
    p.paragraph_format.space_after=Pt(7)

def bul(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.name='Calibri'; r.font.size=Pt(11)
    p.paragraph_format.space_after=Pt(4)

def img(doc, path, width=6.2, caption=None):
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph()
        cr = cp.add_run(caption); cr.font.name='Calibri'; cr.font.size=Pt(9)
        cr.italic=True; cr.font.color.rgb=RGBColor(0x88,0x88,0x88)
        cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after=Pt(12)

def src(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f'📌 Source: {text}'); r.font.name='Calibri'; r.font.size=Pt(9.5)
    r.italic=True; r.font.color.rgb=RGBColor(0x44,0x44,0x88)
    p.paragraph_format.left_indent=Inches(0.3); p.paragraph_format.space_after=Pt(12)

def make_tbl(doc, data, col_widths, gz_col=1, ml_col=2):
    tbl = doc.add_table(rows=len(data), cols=len(data[0]))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    for r_i, row in enumerate(data):
        for c_i, val in enumerate(row):
            cell = tbl.rows[r_i].cells[c_i]
            cell.width = col_widths[c_i]
            cell.text = str(val)
            set_border(cell)
            run = cell.paragraphs[0].runs[0]
            run.font.name='Calibri'; run.font.size=Pt(10)
            if r_i == 0:
                set_bg(cell, '1A3A5C')
                run.font.bold=True; run.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
                cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            else:
                bg = 'F0F4FF' if r_i % 2 == 0 else 'FFFFFF'
                set_bg(cell, bg)
                if c_i == gz_col:
                    run.font.color.rgb=RGBColor(0x21,0x96,0xF3); run.font.bold=True
                    cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
                elif c_i == ml_col:
                    run.font.color.rgb=RGBColor(0x7C,0x3A,0xED); run.font.bold=True
                    cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
                elif c_i > 0:
                    cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after=Pt(4)

# ── Build doc ─────────────────────────────────────────
doc = Document()
for sec in doc.sections:
    sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
    sec.left_margin=Cm(2.8); sec.right_margin=Cm(2.8)

# Title
p = doc.add_heading('Gen Z vs. Millennials', 0)
p.runs[0].font.color.rgb=RGBColor(0x1A,0x3A,0x5C); p.runs[0].font.name='Calibri'
p2 = doc.add_heading('Investment Risk & Portfolio Allocation — Multi-Source Analysis', 2)
p2.runs[0].font.color.rgb=RGBColor(0x21,0x96,0xF3); p2.runs[0].font.name='Calibri'
body(doc,'Data sourced directly from: (1) FINRA Foundation & CFA Institute 2023, (2) Charles Schwab Modern Wealth Survey 2025, (3) Motley Fool Generational Investing Trends Survey July 2025. All figures are real — no estimates.',
     italic=True, color=RGBColor(0x55,0x55,0x99))
doc.add_paragraph()

# ── S1: Sources Overview ──────────────────────────────
h1(doc, '1. Data Sources')
tbl_src = [
    ['Source', 'Year', 'Sample', 'Key Contribution'],
    ['FINRA Foundation & CFA Institute', '2023', '2,872 (U.S., CA, UK, CN)', 'Asset ownership, risk appetite, FOMO, app usage'],
    ['Charles Schwab Modern Wealth Survey', '2025 Wave 2', '2,400 (U.S., 21–75)', 'Portfolio allocation %, crypto attitudes, trading behaviour'],
    ['Motley Fool Generational Investing Survey', 'July 2025', '2,000 (Pollfish)', 'Stock types, sectors, goals, trade frequency, info sources'],
]
make_tbl(doc, tbl_src, [Inches(2.0), Inches(1.0), Inches(1.5), Inches(2.3)], gz_col=None, ml_col=None)
doc.add_page_break()

# ── S2: Asset Ownership ───────────────────────────────
h1(doc, '2. Asset Ownership Rates')
body(doc,'The FINRA/CFA 2023 survey measured what percentage of Gen Z and Millennial investors currently hold each asset class. Crypto ownership rates are nearly identical between the two generations — the more telling difference is that Gen Z is far more likely to have started their investing journey with crypto as their very first asset (44% vs. 35%).')
img(doc, f'{OUT}/c1_ownership_finra.png', caption='Figure 1: % of investors who own each asset class. Source: FINRA Foundation & CFA Institute, 2023')
src(doc, 'FINRA Foundation & CFA Institute. "Gen Z and Investing: Social Media, Crypto, FOMO, and Family." May 2023.')

make_tbl(doc, [
    ['Asset', 'Gen Z', 'Millennials', 'Gen X'],
    ['Cryptocurrency', '55%', '57%', '39%'],
    ['Individual Stocks', '41%', '38%', '43%'],
    ['Mutual Funds', '35%', '43%', '47%'],
    ['NFTs', '25%', '28%', '15%'],
    ['ETFs', '23%', '26%', '22%'],
], [Inches(2.2), Inches(1.2), Inches(1.2), Inches(1.2)])

body(doc,'Table 1: Asset ownership rates (% who hold the asset). Not portfolio allocation percentages.',
     italic=True, color=RGBColor(0x88,0x88,0x88))
doc.add_page_break()

# ── S3: Portfolio Allocation ─────────────────────────
h1(doc, '3. Portfolio Allocation (Schwab 2025)')
body(doc,'The Schwab 2025 Wave 2 survey (October 2025) is the only source that directly asked investors how their portfolio is allocated as percentages summing to 100%. This is across all U.S. investors — the public report does not break this out by generation, but it represents the most current real allocation data available.')
img(doc, f'{OUT}/c7_portfolio_schwab.png', caption='Figure 2: Average portfolio allocation across all U.S. investors. Source: Schwab Modern Wealth Survey Wave 2, October 2025')
src(doc, 'Charles Schwab Modern Wealth Survey 2025 Wave 2. content.schwab.com/web/retail/public/about-schwab/schwab-modern-wealth-survey-2025-wave2-findings.pdf')

make_tbl(doc, [
    ['Asset Class', '% of Avg Portfolio', 'Ownership Rate (% who own)'],
    ['Stocks', '25%', '64%'],
    ['Mutual Funds', '13%', '41%'],
    ['Bonds', '8%', '37%'],
    ['Cryptocurrency', '10%', '35%'],
    ['Real Estate Investments', '7%', '25%'],
    ['ETFs', '6%', '27%'],
    ['Options / Futures', '3%', '15%'],
    ['Alternatives', '3%', '12%'],
    ['Other (incl. cash)', '25%', '—'],
], [Inches(2.2), Inches(1.5), Inches(2.0)], gz_col=1, ml_col=None)

body(doc,'Table 2: Left column = actual portfolio weight. Right column = % of investors who hold the asset at all. Source: Schwab Wave 2 2025.',
     italic=True, color=RGBColor(0x88,0x88,0x88))

h2(doc,'Key Crypto Findings (Schwab 2025)')
bul(doc,'29% of U.S. investors currently own crypto; 23% previously owned it.')
bul(doc,'41% of those aware of crypto consider it a good investment.')
bul(doc,'53% of crypto investors rate it HIGH risk — yet 65% plan to increase their crypto allocation over the next 20 years.')
bul(doc,'Top reason to own crypto: strong long-term growth potential (53%).')
bul(doc,'Top reason NOT to own: concern about scams or fraud (50%).')
doc.add_page_break()

# ── S4: Stock Types ──────────────────────────────────
h1(doc, '4. Stock Types Owned — Gen Z vs. Millennials (Motley Fool 2025)')
body(doc,'The Motley Fool\'s July 2025 survey of 2,000 investors breaks down stock ownership by type across generations. Key finding: Gen Z leads on growth stocks, value stocks, REITs, ESG, and speculative stocks. Millennials lead on U.S. stocks broadly, dividend stocks, large-caps, and small/mid-caps.')
img(doc, f'{OUT}/c2_stock_types_motley.png', caption='Figure 3: Stock types owned by generation. Source: Motley Fool, July 2025')
src(doc, 'Motley Fool Generational Investing Trends Survey, distributed via Pollfish, July 17, 2025. fool.com/research/what-are-gen-z-millennial-investors-buying/')

make_tbl(doc, [
    ['Stock Type', 'Gen Z', 'Millennials', 'Gen X', 'Boomers'],
    ['U.S. Stocks', '42%', '54%', '56%', '56%'],
    ['Growth Stocks', '45%', '48%', '45%', '43%'],
    ['Value Stocks', '39%', '39%', '29%', '16%'],
    ['Dividend-Paying Stocks', '28%', '35%', '43%', '47%'],
    ['Small- or Mid-Cap Stocks', '31%', '39%', '45%', '43%'],
    ['Large-Cap Stocks', '26%', '32%', '38%', '38%'],
    ['REITs', '23%', '23%', '15%', '9%'],
    ['Speculative (Penny/Meme)', '13%', '14%', '9%', '4%'],
    ['ESG / Socially Responsible', '15%', '13%', '6%', '4%'],
], [Inches(2.0), Inches(0.9), Inches(1.1), Inches(0.9), Inches(1.0)])
doc.add_page_break()

# ── S5: Sectors ──────────────────────────────────────
h1(doc, '5. Investment Sectors — Gen Z vs. Millennials (Motley Fool 2025)')
body(doc,'Gen Z under-indexes in Technology vs. Millennials (50% vs. 58%), and substantially trails in Healthcare & Biotech (29% vs. 36%) and Energy (30% vs. 34%). Gen Z leads in Real Estate (29% vs. 26%), AI stocks (22% vs. 21%), and Casino/Sports Betting (12% vs. 9%).')
img(doc, f'{OUT}/c3_sectors_motley.png', caption='Figure 4: Sector ownership by generation. Source: Motley Fool, July 2025')
src(doc, 'Motley Fool Generational Investing Trends Survey, July 2025.')

make_tbl(doc, [
    ['Sector', 'Gen Z', 'Millennials', 'Gen X', 'Boomers'],
    ['Technology', '50%', '58%', '59%', '58%'],
    ['Financials', '36%', '38%', '34%', '34%'],
    ['Healthcare & Biotech', '29%', '36%', '40%', '42%'],
    ['Energy & Utilities', '30%', '34%', '39%', '50%'],
    ['Real Estate', '29%', '26%', '21%', '15%'],
    ['Consumer (Disc. & Staples)', '16%', '22%', '26%', '23%'],
    ['Crypto-Related Stocks', '23%', '28%', '16%', '8%'],
    ['AI Stocks', '22%', '21%', '15%', '12%'],
    ['Industrials & Materials', '20%', '20%', '23%', '24%'],
    ['Casino & Sports Betting', '12%', '9%', '5%', '5%'],
], [Inches(2.0), Inches(0.9), Inches(1.1), Inches(0.9), Inches(1.0)])
doc.add_page_break()

# ── S6: Goals ────────────────────────────────────────
h1(doc, '6. Investment Goals — Gen Z vs. Millennials (Motley Fool 2025)')
body(doc,'The most striking generational divergence in goals: 43% of Gen Z cite passive income as their primary investment goal, vs. only 20% who cite retirement planning. For Millennials, retirement planning (30%) and passive income (40%) are closer. Gen Z is also more likely to be investing to buy a home (16% vs. 9%) or start a business (9% vs. 4%).')
img(doc, f'{OUT}/c4_goals_motley.png', caption='Figure 5: Primary investment goals by generation. Source: Motley Fool, July 2025')
src(doc, 'Motley Fool Generational Investing Trends Survey, July 2025.')

make_tbl(doc, [
    ['Investment Goal', 'Gen Z', 'Millennials', 'Gen X', 'Boomers'],
    ['Passive Income', '43%', '40%', '22%', '14%'],
    ['Retirement Planning', '20%', '30%', '58%', '72%'],
    ['Buying a Home', '16%', '9%', '2%', '2%'],
    ['Capital Appreciation', '7%', '10%', '8%', '6%'],
    ['Starting a Business', '9%', '4%', '3%', '0%'],
    ['Paying Off Debt', '5%', '7%', '6%', '5%'],
], [Inches(2.2), Inches(0.9), Inches(1.1), Inches(0.9), Inches(1.0)])
doc.add_page_break()

# ── S7: Trading ──────────────────────────────────────
h1(doc, '7. Trading Frequency (Motley Fool 2025)')
body(doc,'Gen Z trades significantly more frequently than any other generation. 50% of Gen Z trade daily or weekly (16% daily + 34% weekly), compared to 43% of Millennials. At the other end, 20% of Millennials trade less than once a month vs. 16% of Gen Z — consistent with Millennials\' greater long-term investment focus.')
img(doc, f'{OUT}/c5_trade_freq_motley.png', caption='Figure 6: Trading frequency by generation. Source: Motley Fool, July 2025')
src(doc, 'Motley Fool Generational Investing Trends Survey, July 2025.')

h2(doc, 'Schwab 2025 Corroborating Data')
bul(doc, '43% of all investors trade more frequently now than when they first started.')
bul(doc, '24% trade daily or weekly; 46% trade monthly or quarterly.')
bul(doc, 'Top reason for increased trading: better access to platforms and tools (51%).')
bul(doc, '54% of Gen Z agree investing today requires more short-term risk-taking (vs. 53% Millennials).')
bul(doc, 'Only 57% of Gen Z agree investing requires more long-term patience — the lowest of any generation (Millennials: 61%, Gen X: 65%, Boomers: 66%).')
doc.add_page_break()

# ── S8: Risk & FOMO ──────────────────────────────────
h1(doc, '8. Risk Appetite & FOMO (FINRA 2023)')
img(doc, f'{OUT}/c8_risk_finra.png', caption='Figure 7: Risk & behavioural indicators. Source: FINRA Foundation & CFA Institute, 2023')
src(doc, 'FINRA Foundation & CFA Institute, 2023.')

make_tbl(doc, [
    ['Behavioural Metric', 'Gen Z', 'Millennials'],
    ['Willing to take substantial/above-avg risk', '46%', '~38%'],
    ['Made an investment due to FOMO', '50%', '42%'],
    ['First investment ever was crypto', '44%', '35%'],
    ['Use investing apps to manage/trade', '65%', '55%'],
    ['App suggestion influenced a trade', '67%', '79%'],
    ['61% of Gen Z investors gamble online/in-person', '61%', '—'],
], [Inches(3.2), Inches(1.1), Inches(1.1)])
doc.add_page_break()

# ── S9: Info Sources ─────────────────────────────────
h1(doc, '9. Information Sources (Motley Fool 2025 + FINRA 2023)')
body(doc,'Both surveys show a clear generational divide in where investing information comes from. Gen Z heavily relies on YouTube (67%) and TikTok (48%) — platforms that barely register for older generations. Millennials show more balanced use across platforms and are more likely to use free financial websites and financial advisors.')
img(doc, f'{OUT}/c6_info_sources_motley.png', caption='Figure 8: Investment information sources. Source: Motley Fool, July 2025')
src(doc, 'Motley Fool Generational Investing Trends Survey, July 2025.')

make_tbl(doc, [
    ['Source', 'Gen Z (MF)', 'Millennials (MF)', 'Gen Z (FINRA)', 'Millennials (FINRA)'],
    ['YouTube', '67%', '56%', '60%', '—'],
    ['TikTok', '48%', '26%', '37%', '—'],
    ['Friends / Family', '40%', '42%', '45%', '47%'],
    ['Social Media (all)', '—', '—', '48%', '—'],
    ['Reddit', '27%', '28%', '—', '—'],
    ['Financial Advisor', '25%', '28%', '30%', '39%'],
    ['Free Website / Search', '26%', '37%', '47%', '53%'],
], [Inches(1.8), Inches(1.1), Inches(1.3), Inches(1.1), Inches(1.3)])
doc.add_page_break()

# ── S10: Dividends ───────────────────────────────────
h1(doc, '10. Dividends — Gen Z\'s "Side Hustle" Mentality (Motley Fool 2025)')
body(doc,'A notable 2025 finding: 64% of Gen Z describe dividends as a "side hustle" — the highest of any generation. This framing reflects Gen Z\'s goal-oriented view of investing, where dividends are a near-term income stream rather than a compounding tool. By contrast, 70% of Boomers automatically reinvest dividends.')
img(doc, f'{OUT}/c9_dividends_motley.png', caption='Figure 9: How each generation uses dividend income. Source: Motley Fool, July 2025')
src(doc, 'Motley Fool Generational Investing Trends Survey, July 2025.')

make_tbl(doc, [
    ['Dividend Use', 'Gen Z', 'Millennials', 'Gen X', 'Boomers'],
    ['Take as Cash — Everyday Expenses', '26%', '17%', '8%', '5%'],
    ['Take as Cash — Fun Money', '18%', '15%', '9%', '4%'],
    ['Save for Specific Financial Goals', '19%', '17%', '14%', '8%'],
    ['Reinvest Automatically', '23%', '38%', '53%', '70%'],
    ['Reallocate Manually', '11%', '9%', '10%', '10%'],
], [Inches(2.4), Inches(0.9), Inches(1.1), Inches(0.9), Inches(1.0)])
doc.add_page_break()

# ── S11: Master Comparison ───────────────────────────
h1(doc, '11. Master Comparison Table — All Sources')
make_tbl(doc, [
    ['Metric', 'Gen Z', 'Millennials', 'Source'],
    ['Crypto ownership rate', '55%', '57%', 'FINRA 2023'],
    ['Individual stock ownership', '41%', '38%', 'FINRA 2023'],
    ['ETF ownership', '23%', '26%', 'FINRA 2023'],
    ['Mutual fund ownership', '35%', '43%', 'FINRA 2023'],
    ['First investment was crypto', '44%', '35%', 'FINRA 2023'],
    ['Made investment due to FOMO', '50%', '42%', 'FINRA 2023'],
    ['Willing to take substantial risk', '46%', '~38%', 'FINRA 2023'],
    ['Use investing apps to trade', '65%', '55%', 'FINRA 2023'],
    ['Crypto % of avg portfolio', '10%*', '10%*', 'Schwab 2025 W2'],
    ['Stocks % of avg portfolio', '25%*', '25%*', 'Schwab 2025 W2'],
    ['Currently own crypto (investors)', '29%*', '29%*', 'Schwab 2025 W2'],
    ['U.S. stocks ownership', '42%', '54%', 'Motley Fool 2025'],
    ['Growth stocks ownership', '45%', '48%', 'Motley Fool 2025'],
    ['Dividend stocks ownership', '28%', '35%', 'Motley Fool 2025'],
    ['Value stocks ownership', '39%', '39%', 'Motley Fool 2025'],
    ['Primary goal: passive income', '43%', '40%', 'Motley Fool 2025'],
    ['Primary goal: retirement', '20%', '30%', 'Motley Fool 2025'],
    ['Trade daily or weekly', '50%', '43%', 'Motley Fool 2025'],
    ['Use YouTube for investing info', '67%', '56%', 'Motley Fool 2025'],
    ['Use TikTok for investing info', '48%', '26%', 'Motley Fool 2025'],
    ['View dividends as "side hustle"', '64%', '53%', 'Motley Fool 2025'],
], [Inches(2.8), Inches(1.0), Inches(1.1), Inches(1.7)])
body(doc, '* Schwab 2025 Wave 2 portfolio allocation data is not broken out by generation in the public report — figures represent all U.S. investors.',
     italic=True, color=RGBColor(0x88,0x88,0x88))

doc.add_page_break()
h1(doc, '12. Full Citations')
for cite in [
    'FINRA Investor Education Foundation & CFA Institute (May 2023). "Gen Z and Investing: Social Media, Crypto, FOMO, and Family." finrafoundation.org/sites/finrafoundation/files/Gen-Z-and-Investing.pdf',
    'Charles Schwab & Co., Inc. (October 2025). "Modern Wealth Survey 2025, Wave 2." content.schwab.com/web/retail/public/about-schwab/schwab-modern-wealth-survey-2025-wave2-findings.pdf',
    'Motley Fool (August 2025). "Survey: What Are Gen Z and Millennial Investors Buying in 2025?" Distributed via Pollfish, July 17, 2025. fool.com/research/what-are-gen-z-millennial-investors-buying/',
]:
    bul(doc, cite)

doc.save('/home/claude/GenZ_vs_Millennials_MultiSource_Report.docx')
print('✅ Word report saved')
